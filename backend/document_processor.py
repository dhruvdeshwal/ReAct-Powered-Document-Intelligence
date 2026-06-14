"""
document_processor.py
----------------------
Handles ingestion of PDF, DOCX, and TXT files:
- Extracts raw text (with page/paragraph metadata where possible)
- Cleans text (removes excess whitespace)
- Splits into overlapping chunks suitable for embedding
"""

import os
from typing import List, Dict
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 800))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 100))


def _clean_text(text: str) -> str:
    """Collapse excessive whitespace/newlines while preserving paragraph breaks."""
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]  # drop empty lines
    return "\n".join(lines)


def load_pdf(file_path: str) -> List[Dict]:
    """Extract text from a PDF, one record per page."""
    records = []
    reader = PdfReader(file_path)
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = _clean_text(text)
        if text:
            records.append({"text": text, "page": page_num})
    return records


def load_docx(file_path: str) -> List[Dict]:
    """Extract text from a DOCX file. DOCX has no native 'pages', so treat as one block,
    but also capture tables as separate records for better retrieval."""
    records = []
    doc = DocxDocument(file_path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        records.append({"text": _clean_text("\n".join(paragraphs)), "page": None})

    for t_idx, table in enumerate(doc.tables, start=1):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        table_text = _clean_text("\n".join(rows_text))
        if table_text:
            records.append({"text": f"[Table {t_idx}]\n{table_text}", "page": None})

    return records


def load_txt(file_path: str) -> List[Dict]:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    text = _clean_text(text)
    return [{"text": text, "page": None}] if text else []


def load_document(file_path: str) -> List[Dict]:
    """Dispatch based on file extension. Returns a list of {text, page} records."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def chunk_document(file_path: str) -> List[Dict]:
    """
    Load a document and split it into chunks ready for embedding.

    Returns a list of dicts:
        {
            "text": "<chunk text>",
            "metadata": {
                "source": "<filename>",
                "page": <page number or None>,
                "chunk_index": <int>
            }
        }
    """
    filename = os.path.basename(file_path)
    records = load_document(file_path)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    chunks = []
    chunk_index = 0
    for record in records:
        for piece in splitter.split_text(record["text"]):
            chunks.append({
                "text": piece,
                "metadata": {
                    "source": filename,
                    "page": record.get("page"),
                    "chunk_index": chunk_index,
                }
            })
            chunk_index += 1

    return chunks


if __name__ == "__main__":
    # Quick standalone test
    import sys
    if len(sys.argv) < 2:
        print("Usage: python document_processor.py <file_path>")
    else:
        result = chunk_document(sys.argv[1])
        print(f"Generated {len(result)} chunks")
        for c in result[:3]:
            print("---")
            print(c["metadata"])
            print(c["text"][:200])